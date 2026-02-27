"""DOCX document parser using python-docx."""
from io import BytesIO
from typing import Dict, Any
from zipfile import BadZipFile, ZipFile

from smf.parsers.base import DocumentParser


class DOCXParser(DocumentParser):
    """Parser for DOCX documents using python-docx library."""

    def __init__(self):
        """Initialize the DOCX parser."""
        try:
            from docx import Document
            self.Document = Document
            self._available = True
        except ImportError:
            self._available = False

    def parse(self, content: bytes) -> str:
        """Parse DOCX content and extract text.
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Extracted text from DOCX
            
        Raises:
            Exception: If python-docx is not installed, file contains macros, or parsing fails
        """
        if not self._available:
            raise Exception("python-docx not installed. Install with: pip install python-docx")

        try:
            # Use a single BytesIO object for both macro checking and text extraction
            docx_stream = BytesIO(content)

            # Check for macros (DOCM files have vbaProject.bin)
            try:
                with ZipFile(docx_stream) as docx_zip:
                    if "word/vbaProject.bin" in docx_zip.namelist():
                        raise Exception("DOCX file contains macros and is not allowed for security reasons")
            except BadZipFile:
                raise Exception("Invalid DOCX file format")

            # Seek back to the start for Document()
            docx_stream.seek(0)
            doc = self.Document(docx_stream)
            text_parts = [para.text for para in doc.paragraphs]

            # Extract text from tables
            # Note: This iterates through all cells which is simple but may include duplicates
            # for merged cells. For better efficiency with large documents, consider using
            # doc.element.xpath() for direct cell access or caching cell references.
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

            return "\n".join(text_parts)
        except Exception as exc:
            if "macros" in str(exc):
                raise  # Re-raise macro security exceptions
            raise Exception(f"Failed to extract DOCX text: {exc}") from exc

    def supports_format(self, filename: str) -> bool:
        """Check if this parser supports DOCX files.
        
        Args:
            filename: Name of the file to check
            
        Returns:
            True if filename ends with .docx
        """
        return filename.lower().endswith('.docx')

    def get_metadata(self, content: bytes) -> Dict[str, Any]:
        """Extract metadata from DOCX.
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Dictionary with DOCX metadata (paragraphs, tables, etc.)
        """
        if not self._available:
            return {}

        try:
            docx_stream = BytesIO(content)
            docx_stream.seek(0)
            doc = self.Document(docx_stream)
            
            return {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }
        except Exception:
            return {}
