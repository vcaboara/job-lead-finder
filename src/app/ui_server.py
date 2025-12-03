"""FastAPI UI server for job lead finder.

Reconstructed after refactor: serves HTML template, search API,
configuration endpoints, and leads retrieval.
"""

import json
import logging
import os
import re
import time
from io import BytesIO
from pathlib import Path
from typing import Any, Dict
from zipfile import ZipFile, BadZipFile

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, PlainTextResponse
from pydantic import BaseModel

from .config_manager import get_search_preferences, load_config, save_config, scan_entity, scan_instructions, validate_url
from .job_finder import generate_job_leads, save_to_file
from .job_tracker import STATUS_NEW, VALID_STATUSES, get_tracker
from .link_validator import validate_link

# Optional imports for PDF/DOCX support
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s")

app = FastAPI(title="Job Lead Finder", version="0.1.1")

LEADS_FILE = Path("leads.json")
RESUME_FILE = Path("resume.txt")
UPLOADS_DIR = Path("uploads")
UPLOADS_DIR.mkdir(exist_ok=True)

# Progress tracking for search operations
search_progress: Dict[str, dict] = {}


class SearchRequest(BaseModel):
    query: str
    resume: str | None = None
    count: int = 5
    model: str | None = None
    evaluate: bool = False
    min_score: int = 60  # Minimum score threshold (0-100 scale) for filtering results


class HealthResponse(BaseModel):
    status: str
    api_key_configured: bool


class SystemInstructionsRequest(BaseModel):
    instructions: str


class BlockedEntity(BaseModel):
    type: str
    value: str


class ConfigResponse(BaseModel):
    system_instructions: str
    blocked_entities: list[BlockedEntity]
    region: str


class BlockedEntityRequest(BaseModel):
    entity: str
    entity_type: str  # 'site' or 'employer'


class ValidateLinkRequest(BaseModel):
    url: str


@app.get("/", response_class=HTMLResponse)
def index():
    """Serve the main web UI interface.
    
    Returns:
        HTMLResponse: The main application HTML page.
        
    Raises:
        HTTPException: 500 if template file cannot be read.
    """
    html_path = Path(__file__).parent / "templates" / "index.html"
    try:
        return HTMLResponse(html_path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise HTTPException(status_code=500, detail="UI template not found") from exc


@app.get("/health", response_model=HealthResponse)
def health():
    api_key = os.getenv("GEMINI_API_KEY", "")
    return HealthResponse(status="healthy", api_key_configured=bool(api_key))


@app.get("/api/changelog")
def get_changelog():
    """Get the project changelog documenting all version changes.
    
    Returns:
        Plain text changelog following Keep a Changelog format.
        Includes version history, API changes, and feature documentation.
        
    Raises:
        HTTPException: 404 if CHANGELOG.md not found.
    """
    changelog_path = Path(__file__).parent.parent.parent / "CHANGELOG.md"
    try:
        return PlainTextResponse(changelog_path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail="Changelog not found") from exc


@app.get("/api/config", response_model=ConfigResponse)
def get_config():
    cfg = load_config()
    return ConfigResponse(
        system_instructions=cfg.get("system_instructions", ""),
        blocked_entities=cfg.get("blocked_entities", []),
        region=cfg.get("region", ""),
    )


@app.post("/api/config/system-instructions", response_model=ConfigResponse)
def update_system_instructions(req: SystemInstructionsRequest):
    findings = scan_instructions(req.instructions)
    if findings:
        raise HTTPException(status_code=400, detail={"error": "Rejected by scanner", "findings": findings})
    cfg = load_config()
    cfg["system_instructions"] = req.instructions
    save_config(cfg)
    return ConfigResponse(
        system_instructions=cfg.get("system_instructions", ""),
        blocked_entities=cfg.get("blocked_entities", []),
        region=cfg.get("region", ""),
    )


@app.post("/api/search")
def search(req: SearchRequest):
    """Search for job leads with timeout protection.
    
    Implements a maximum search timeout of 5 minutes to prevent indefinite hangs.
    Progress can be monitored via /api/search/progress/{search_id} endpoint.
    """
    # Allow search without an API key; provider will fallback internally
    # Load resume from file if not provided in request
    resume_text = req.resume
    if not resume_text and RESUME_FILE.exists():
        resume_text = RESUME_FILE.read_text(encoding="utf-8")

    # Auto-evaluate if resume is available
    should_evaluate = req.evaluate or bool(resume_text)

    # Create unique search ID for progress tracking
    search_id = f"search_{int(time.time() * 1000)}"
    start_time = time.time()
    search_progress[search_id] = {
        "search_id": search_id,
        "status": "starting",
        "message": f"Requesting {req.count} job listings...",
        "attempt": 0,
        "valid_count": 0,
        "timestamp": start_time,
        "start_time": start_time,
    }

    logger.info("[%s] Search started: query='%s', count=%d, model=%s", search_id, req.query, req.count, req.model)

    # Maximum search timeout: 5 minutes
    MAX_SEARCH_TIMEOUT = 300  # seconds

    try:
        # Request 10x more jobs to account for filtering (oversample strategy)
        # This helps ensure we get the requested count after validation
        # With high invalid link rates (soft 404s, hallucinations), we need aggressive oversampling
        search_prefs = get_search_preferences()
        oversample_multiplier = search_prefs.get("oversample_multiplier", 10)
        initial_request_count = req.count * oversample_multiplier
        max_retries = 1  # Reduced from 2 for faster response
        all_valid_leads = []
        seen_links = set()  # Track unique job links to avoid duplicates

        logger.info(
            "[%s] Configuration: oversample=%dx, initial_request=%d, max_retries=%d",
            search_id,
            oversample_multiplier,
            initial_request_count,
            max_retries
        )

        for attempt in range(max_retries + 1):
            # Check if we've exceeded maximum search time
            elapsed_time = time.time() - start_time
            if elapsed_time > MAX_SEARCH_TIMEOUT:
                logger.warning(
                    "[%s] Search timeout after %.1fs (max: %ds) - returning %d jobs found so far",
                    search_id,
                    elapsed_time,
                    MAX_SEARCH_TIMEOUT,
                    len(all_valid_leads),
                )
                search_progress[search_id].update(
                    {
                        "status": "timeout",
                        "message": f"Search timeout after {elapsed_time:.0f}s - returning {len(all_valid_leads)} jobs",
                    }
                )
                break

            # Calculate how many more we need
            needed = req.count - len(all_valid_leads)
            if needed <= 0:
                break

            # Request extra to account for filtering
            request_count = needed * oversample_multiplier if attempt > 0 else initial_request_count

            # Update progress
            attempt_start = time.time()
            elapsed = attempt_start - start_time
            search_progress[search_id].update(
                {
                    "status": "fetching",
                    "attempt": attempt + 1,
                    "message": (
                        f"Attempt {attempt + 1}/{max_retries + 1}: "
                        f"Searching {request_count} jobs across providers... (elapsed: {elapsed:.1f}s)"
                    ),
                }
            )

            logger.info(
                "[%s] Attempt %d/%d: requesting %d jobs (need %d more valid), elapsed=%.1fs",
                search_id,
                attempt + 1,
                max_retries + 1,
                request_count,
                needed,
                elapsed
            )

            raw_leads = generate_job_leads(
                query=req.query,
                resume_text=resume_text or "No resume provided.",
                count=request_count,
                model=req.model,
                evaluate=should_evaluate,
                verbose=False,
            )

            fetch_elapsed = time.time() - attempt_start
            logger.info("[%s] Fetched %d raw jobs in %.1fs", search_id, len(raw_leads), fetch_elapsed)

            # Update progress with provider stats
            total_elapsed = time.time() - start_time
            search_progress[search_id].update(
                {
                    "status": "filtering",
                    "message": (
                        f"Fetched {len(raw_leads)} jobs in {fetch_elapsed:.1f}s. "
                        f"Validating links... (total: {total_elapsed:.1f}s)"
                    ),
                }
            )

            # Process and filter leads
            filter_start = time.time()
            valid_leads = _process_and_filter_leads(raw_leads)
            logger.info(
                "[%s] Filtered to %d valid jobs (removed %d invalid)",
                search_id,
                len(valid_leads),
                len(raw_leads) - len(valid_leads)
            )

            # Deduplicate by link and filter out hidden jobs
            tracker = get_tracker()
            hidden_count = 0
            duplicate_count = 0
            for lead in valid_leads:
                link = lead.get("link", "")
                if link and link not in seen_links:
                    # Check if job is hidden
                    if not tracker.is_job_hidden(lead):
                        seen_links.add(link)
                        all_valid_leads.append(lead)
                        # Track new job automatically as "new" status
                        tracker.track_job(lead, STATUS_NEW)
                    else:
                        hidden_count += 1
                elif link:
                    duplicate_count += 1

            filter_elapsed = time.time() - filter_start
            logger.info(
                "[%s] After deduplication: %d unique jobs (filtered %d hidden, %d duplicates) in %.1fs",
                search_id,
                len(all_valid_leads),
                hidden_count,
                duplicate_count,
                filter_elapsed
            )

            total_elapsed = time.time() - start_time
            search_progress[search_id].update(
                {
                    "valid_count": len(all_valid_leads),
                    "message": f"Found {len(all_valid_leads)} valid jobs so far... (total: {total_elapsed:.1f}s)",
                }
            )

            # Stop if we have enough or got no new results
            if len(all_valid_leads) >= req.count or len(valid_leads) == 0:
                break

        # Take only the requested count
        final_leads = all_valid_leads[: req.count]

        # Enrich leads with tracking data (status, notes, etc.)
        tracker = get_tracker()
        for lead in final_leads:
            from .job_tracker import generate_job_id

            job_id = generate_job_id(lead)
            tracked_job = tracker.get_job(job_id)
            if tracked_job:
                lead["job_id"] = job_id
                lead["tracking_status"] = tracked_job.get("status", STATUS_NEW)
                lead["tracking_notes"] = tracked_job.get("notes", "")
                lead["company_link"] = tracked_job.get("company_link")

        # Filter by minimum score if evaluation was performed AND jobs have scores
        # Only apply filter if at least some jobs were actually scored
        jobs_with_scores = [lead for lead in final_leads if lead.get("score") is not None]
        if should_evaluate and req.min_score > 0 and len(jobs_with_scores) > 0:
            before_filter = len(final_leads)
            final_leads = [
                lead for lead in final_leads if lead.get("score") is not None and lead.get("score", 0) >= req.min_score
            ]
            filtered_count = before_filter - len(final_leads)
            if filtered_count > 0:
                logger.info("[%s] Filtered out %d jobs below score threshold %d", search_id, filtered_count, req.min_score)
                if len(final_leads) == 0:
                    logger.warning(
                        "[%s] All %d jobs filtered out by min_score=%d. Consider lowering the score threshold.",
                        search_id, before_filter, req.min_score
                    )
        elif should_evaluate and req.min_score > 0 and len(jobs_with_scores) == 0:
            logger.warning(
                "[%s] Score filter requested but no jobs have scores - showing all %d jobs",
                search_id,
                len(final_leads)
            )

        total_elapsed = time.time() - start_time
        logger.info("[%s] Search complete: %d jobs returned in %.1fs", search_id, len(final_leads), total_elapsed)

        search_progress[search_id].update(
            {
                "status": "complete",
                "message": f"Search complete: {len(final_leads)} jobs found in {total_elapsed:.1f}s",
                "valid_count": len(final_leads),
                "elapsed": total_elapsed,
            }
        )

        save_to_file(final_leads, str(LEADS_FILE))
        return JSONResponse(
            {
                "status": "success",
                "query": req.query,
                "count": len(final_leads),
                "leads": final_leads,
                "filtered_count": len(all_valid_leads) - len(final_leads),
                "total_fetched": len(all_valid_leads),
                "search_id": search_id,
            }
        )
    except Exception as e:
        search_progress[search_id].update({"status": "error", "message": f"Error: {str(e)}"})
        raise HTTPException(status_code=500, detail=str(e)) from e


@app.get("/api/search/progress/{search_id}")
def get_search_progress(search_id: str):
    """Get real-time progress updates for an active search.
    
    Returns current status, message, elapsed time, and job count for a search.
    Useful for polling during long-running searches.
    
    Args:
        search_id: Unique search identifier from /api/search response
        
    Returns:
        Progress object with status, message, elapsed time, and valid_count
    """
    if search_id not in search_progress:
        raise HTTPException(status_code=404, detail=f"Search {search_id} not found")
    
    progress = search_progress[search_id].copy()
    # Calculate elapsed time if search is still running
    if "start_time" in progress and progress.get("status") not in ["complete", "error", "timeout"]:
        progress["elapsed"] = time.time() - progress["start_time"]
    
    return JSONResponse(progress)


def _process_and_filter_leads(raw_leads: list) -> list:
    """Process raw leads and filter out blocked/invalid ones."""
    # Load config for blocked entities
    cfg = load_config()
    blocked = cfg.get("blocked_entities", [])
    blocked_sites = {e.get("value", "").lower() for e in blocked if e.get("type") == "site"}
    blocked_employers = {e.get("value", "").lower() for e in blocked if e.get("type") == "employer"}

    def is_blocked(lead_link: str, company: str) -> bool:
        from urllib.parse import urlparse

        # Company check
        if company and company.lower() in blocked_employers:
            return True
        # Site/domain check
        if not lead_link:
            return False
        try:
            parsed = urlparse(lead_link)
            host = parsed.netloc.lower()
            host = host[4:] if host.startswith("www.") else host
            # Direct match or suffix match (e.g., sub.domain.com endswith domain.com)
            for site in blocked_sites:
                if host == site or host.endswith("." + site):
                    return True
        except Exception:
            return False
        return False

    # Parallelize link validation for speed
    import concurrent.futures
    from urllib.parse import urlparse

    # Pre-filter blocked leads
    unblocked_leads = [lead for lead in raw_leads if not is_blocked(lead.get("link", ""), lead.get("company", ""))]
    print(f"_process_and_filter_leads: {len(raw_leads)} raw leads, {len(unblocked_leads)} after blocking filter")

    # Validate all links in parallel (5-10x faster than sequential)
    links_to_validate = [lead.get("link", "") for lead in unblocked_leads]
    validated_results = {}

    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
        future_to_link = {
            executor.submit(validate_link, link, timeout=10, verbose=False): link for link in links_to_validate if link
        }
        for future in concurrent.futures.as_completed(future_to_link):
            link = future_to_link[future]
            try:
                validated_results[link] = future.result()
            except Exception as e:
                print(f"Link validation exception for {link}: {e}")
                validated_results[link] = {"valid": False, "status_code": None, "error": "validation_failed"}

    processed_leads = []
    filtered_reasons = {}
    for lead in unblocked_leads:
        link = lead.get("link", "")
        link_info = validated_results.get(link) if link else {"valid": False, "status_code": None, "error": "no_link"}
        # Exclude bad links: 403, 404, localhost/127.0.0.1, search-result pages, and generic career pages
        parsed = urlparse(link) if link else None
        host = parsed.netloc.lower() if parsed else ""
        host = host[4:] if host.startswith("www.") else host
        is_local = host in {"localhost", "127.0.0.1"}
        path = parsed.path.lower() if parsed else ""
        query = parsed.query.lower() if parsed else ""

        # Detect generic career/jobs pages (not specific job postings)
        # Allow job board sites (LinkedIn, Indeed, Glassdoor, etc.)
        host_is_job_board = any(
            job_board in host
            for job_board in [
                "linkedin.com",
                "indeed.com",
                "glassdoor",
                "github.com",
                "remote",
                "workable.com",
                "greenhouse.io",
                "lever.co",
                "jobvite.com",
                "applytojob.com",
            ]
        )

        generic_patterns = [
            "/careers",
            "/career",
            "/employment",
            "/opportunities",
            "/join-us",
            "/work-with-us",
        ]
        # Only check generic patterns if NOT on a job board site
        is_generic_page = False
        if not host_is_job_board:
            is_generic_page = any(
                path.rstrip("/") == pattern or path.rstrip("/") + "/" == pattern + "/" for pattern in generic_patterns
            )

        # Be more lenient with "search" pages on job boards - they often work
        looks_like_search = False
        if not host_is_job_board:
            looks_like_search = ("/search" in path) or ("jobs/search" in path) or ("q=" in query)

        status = link_info.get("status_code")
        is_excluded_status = status == 404  # Only exclude 404s, not 403 (soft-valid)

        # Track filtering reasons
        filter_reason = None
        if not link_info.get("valid"):
            # Only filter truly broken links - exclude timeouts and generic errors
            error = link_info.get("error") or ""
            if "timeout" not in error.lower() and status == 404:
                filter_reason = f"invalid_link:{error}"
        elif is_excluded_status:
            filter_reason = f"excluded_status:{status}"
        elif is_local:
            filter_reason = "localhost"
        elif looks_like_search:
            filter_reason = "search_page"
        elif is_generic_page:
            filter_reason = "generic_page"

        if filter_reason:
            filtered_reasons[filter_reason] = filtered_reasons.get(filter_reason, 0) + 1
            continue

        lead["link_status_code"] = status
        lead["link_valid"] = True
        lead["link_warning"] = link_info.get("warning")
        lead["link_error"] = link_info.get("error")
        processed_leads.append(lead)

    print(f"_process_and_filter_leads: {len(processed_leads)} passed validation. Filtered: {filtered_reasons}")

    return processed_leads


@app.get("/api/leads")
def get_leads():
    if not LEADS_FILE.exists():
        return JSONResponse({"leads": []})
    try:
        with open(LEADS_FILE, "r", encoding="utf-8") as fh:
            leads = json.load(fh)
        return JSONResponse({"leads": leads})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading leads: {e}") from e


@app.post("/api/config/block-entity", response_model=ConfigResponse)
def add_blocked_entity(req: BlockedEntityRequest):
    """Add a site or employer to the block list."""
    entity = req.entity.strip()
    entity_type = req.entity_type.lower()

    if entity_type not in ("site", "employer"):
        raise HTTPException(status_code=400, detail="entity_type must be 'site' or 'employer'")

    # Validate site domain format
    if entity_type == "site" and not validate_url(entity):
        raise HTTPException(status_code=400, detail="Invalid site domain format")

    # Scan for injection
    findings = scan_entity(entity)
    if findings:
        raise HTTPException(status_code=400, detail={"error": "Rejected by scanner", "findings": findings})

    cfg = load_config()
    blocked = cfg.get("blocked_entities", [])

    # Store as dict with type
    entry = {"type": entity_type, "value": entity}
    if entry not in blocked:
        blocked.append(entry)
        cfg["blocked_entities"] = blocked
        save_config(cfg)

    return ConfigResponse(
        system_instructions=cfg.get("system_instructions", ""),
        blocked_entities=cfg.get("blocked_entities", []),
        region=cfg.get("region", ""),
    )


@app.delete("/api/config/block-entity/{entity_type}/{entity}", response_model=ConfigResponse)
def remove_blocked_entity(entity_type: str, entity: str):
    """Remove a site or employer from the block list."""
    entity = entity.strip()
    entity_type = entity_type.lower()

    cfg = load_config()
    blocked = cfg.get("blocked_entities", [])
    entry = {"type": entity_type, "value": entity}

    if entry in blocked:
        blocked.remove(entry)
        cfg["blocked_entities"] = blocked
        save_config(cfg)

    return ConfigResponse(
        system_instructions=cfg.get("system_instructions", ""),
        blocked_entities=cfg.get("blocked_entities", []),
        region=cfg.get("region", ""),
    )


@app.post("/api/validate-link")
def check_link(req: ValidateLinkRequest):
    """Validate a link is accessible."""
    result = validate_link(req.url, verbose=False)
    return JSONResponse(
        {
            "url": result.get("url", req.url),
            "valid": result.get("valid", False),
            "status_code": result.get("status_code"),
            "error": result.get("error"),
        }
    )


@app.post("/api/upload/resume")
async def upload_resume(file: UploadFile = File(...)):
    """Upload resume file for job matching.
    
    Supports: .txt, .md, .pdf, .docx
    Max size: 5MB
    Security: Scans for injection patterns, macros, and malicious content
    """
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
    ALLOWED_EXTENSIONS = (".txt", ".md", ".pdf", ".docx")
    
    # Validate file size
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400, 
            detail=f"File too large (max {MAX_FILE_SIZE // (1024*1024)}MB, got {len(content) // (1024*1024)}MB)"
        )

    # Validate file type
    if not file.filename or not file.filename.lower().endswith(ALLOWED_EXTENSIONS):
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    # Extract text based on file type
    try:
        if file.filename.lower().endswith(".pdf"):
            text = _extract_pdf_text(content)
        elif file.filename.lower().endswith(".docx"):
            text = _extract_docx_text(content)
        else:  # .txt or .md
            text = content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=400, detail="File must be valid UTF-8 text") from exc
    except Exception as exc:
        file_ext = Path(file.filename).suffix[1:].upper() if Path(file.filename).suffix else 'UNKNOWN'
        raise HTTPException(
            status_code=400, 
            detail=f"Failed to extract text from {file_ext} file: {str(exc)}"
        ) from exc

    # Enhanced security checks
    if not text or not text.strip():
        raise HTTPException(status_code=400, detail="File appears to be empty or contains no extractable text")
    
    if len(text) > 500_000:  # 500KB text limit
        raise HTTPException(
            status_code=400, 
            detail=f"Extracted text too large ({len(text)} chars, max 500,000 chars)"
        )

    # Check for malicious content
    malicious_findings = _check_malicious_content(text)
    if malicious_findings:
        raise HTTPException(
            status_code=400, 
            detail={"error": "File rejected due to security concerns", "findings": malicious_findings}
        )

    # Scan for injection patterns (first 5000 chars)
    findings = scan_instructions(text[:5000])
    if findings:
        raise HTTPException(
            status_code=400, 
            detail={"error": "Rejected by scanner", "findings": findings}
        )

    # Save to resume.txt
    RESUME_FILE.write_text(text, encoding="utf-8")

    return JSONResponse({
        "message": "Resume uploaded successfully", 
        "resume": text, 
        "filename": file.filename,
        "size_bytes": len(content),
        "text_length": len(text)
    })


def _extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF file.
    
    Args:
        content: PDF file bytes
        
    Returns:
        Extracted text from PDF
        
    Raises:
        Exception: If PDF extraction fails
    """
    if not PYPDF_AVAILABLE:
        raise Exception("pypdf not installed. Install with: pip install pypdf")
    
    try:
        pdf = PdfReader(BytesIO(content))
        text_parts = []
        
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        raw_text = "\n".join(text_parts)
        
        # Clean up the extracted text
        # Fix multiple spaces between words
        cleaned_text = re.sub(r'  +', ' ', raw_text)
        # Fix common mojibake (UTF-8 mis-decoded as Windows-1252)
        cleaned_text = cleaned_text.replace('â€"', '—')    # em dash
        cleaned_text = cleaned_text.replace('â€"', '–')    # en dash
        cleaned_text = cleaned_text.replace('â€™', "'")    # apostrophe
        cleaned_text = cleaned_text.replace('â€œ', '"')    # left double quote
        cleaned_text = cleaned_text.replace('â€', '"')    # right double quote
        cleaned_text = cleaned_text.replace('â€¢', '•')    # bullet
        # Normalize line breaks
        cleaned_text = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned_text)
        
        return cleaned_text.strip()
    except Exception as exc:
        raise Exception(f"Failed to extract PDF text: {exc}") from exc


def _extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX file.
    
    Args:
        content: DOCX file bytes
        
    Returns:
        Extracted text from DOCX
        
    Raises:
        Exception: If DOCX extraction fails or contains macros
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    
    try:
        # Use a single BytesIO object for both macro checking and text extraction
        docx_stream = BytesIO(content)
        
        # Check for macros (DOCM files have vbaProject.bin)
        try:
            with ZipFile(docx_stream) as docx_zip:
                if "word/vbaProject.bin" in docx_zip.namelist():
                    raise Exception("DOCX file contains macros and is not allowed for security reasons")
        except BadZipFile:
            raise Exception("Invalid DOCX file format")
        
        # Seek back to the start for Document()
        docx_stream.seek(0)
        doc = Document(docx_stream)
        text_parts = [para.text for para in doc.paragraphs]
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    except Exception as exc:
        if "macros" in str(exc):
            raise  # Re-raise macro security exceptions
        raise Exception(f"Failed to extract DOCX text: {exc}") from exc


def _check_malicious_content(text: str) -> list[str]:
    """Check for malicious content in uploaded file.
    
    Args:
        text: Extracted text content
        
    Returns:
        List of security findings (empty if safe)
    """
    findings = []
    
    # Check for embedded scripts
    script_patterns = [
        "<script", "</script>",
        "javascript:", "vbscript:",
        "onclick=", "onerror=", "onload=", "onmouseover=", "onfocus=",
        "<iframe", "<embed", "<object",
        "eval(", "exec(",
    ]
    
    text_lower = text.lower()
    for pattern in script_patterns:
        if pattern.lower() in text_lower:
            findings.append(f"Suspicious pattern detected: '{pattern}'")
    
    # Check for excessive special characters (possible obfuscation)
    # Exclude common resume punctuation from special character count
    common_punct = set(".,:;()-[]{}•*'\"/\\&+|_")
    special_char_count = sum(1 for c in text if not c.isalnum() and not c.isspace() and c not in common_punct)
    if len(text) > 100 and special_char_count / len(text) > 0.45:
        findings.append(f"Excessive special characters detected ({special_char_count}/{len(text)} = {special_char_count/len(text):.1%})")
    
    # Check for null bytes (binary content)
    if "\x00" in text:
        findings.append("Binary content detected (null bytes found)")
    
    # Check for very long lines (possible attack vector)
    lines = text.split("\n")
    max_line_length = max(len(line) for line in lines) if lines else 0
    if max_line_length > 10000:
        findings.append(f"Extremely long line detected ({max_line_length} chars)")
    
    return findings


@app.get("/api/resume")
def get_resume():
    """Get current resume text."""
    if not RESUME_FILE.exists():
        return JSONResponse({"resume": None})
    try:
        text = RESUME_FILE.read_text(encoding="utf-8")
        return JSONResponse({"resume": text})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading resume: {e}") from e


@app.delete("/api/resume")
def delete_resume():
    """Delete resume file."""
    if not RESUME_FILE.exists():
        raise HTTPException(status_code=404, detail="Resume not found")
    RESUME_FILE.unlink()
    return JSONResponse({"message": "Resume deleted"})


# ============================================================================
# Job Search Configuration Endpoints
# ============================================================================


@app.get("/api/job-config")
def get_job_config():
    """Get job search configuration (location, providers, search params).
    
    Returns:
        JSONResponse: Complete configuration including providers, location, and search settings.
    """
    config = load_config()
    return JSONResponse(config)


@app.post("/api/job-config/location")
def update_location_config(
    default_location: str | None = None,
    prefer_remote: bool | None = None,
    allow_hybrid: bool | None = None,
    allow_onsite: bool | None = None,
):
    """Update location and remote/onsite preferences."""
    from .config_manager import update_location_preferences

    success = update_location_preferences(
        default_location=default_location,
        prefer_remote=prefer_remote,
        allow_hybrid=allow_hybrid,
        allow_onsite=allow_onsite,
    )
    if success:
        return JSONResponse({"message": "Location preferences updated"})
    raise HTTPException(status_code=500, detail="Failed to update location preferences")


@app.post("/api/job-config/provider/{provider_key}")
def toggle_provider(provider_key: str, enabled: bool):
    """Enable or disable a job search provider."""
    from .config_manager import update_provider_status

    success = update_provider_status(provider_key, enabled)
    if success:
        return JSONResponse({"message": f"Provider {provider_key} {'enabled' if enabled else 'disabled'}"})
    raise HTTPException(status_code=404, detail=f"Provider {provider_key} not found")


@app.post("/api/job-config/search")
def update_search_config(req: Dict[str, Any]):
    """Update search parameters."""
    from .config_manager import update_search_preferences

    success = update_search_preferences(
        default_count=req.get("default_count"),
        oversample_multiplier=req.get("oversample_multiplier"),
        enable_ai_ranking=req.get("enable_ai_ranking"),
    )
    if success:
        return JSONResponse({"message": "Search preferences updated"})
    raise HTTPException(status_code=500, detail="Failed to update search preferences")


@app.get("/api/industry-profiles")
def get_industry_profiles():
    """Get list of available industry profiles."""
    from .industry_profiles import list_profiles

    return JSONResponse({"profiles": list_profiles()})


@app.get("/api/industry-profile")
def get_current_industry_profile():
    """Get current industry profile."""
    from .config_manager import get_industry_profile
    from .industry_profiles import get_profile

    current = get_industry_profile()
    profile = get_profile(current)
    return JSONResponse({"current": current, "profile": profile})


@app.post("/api/industry-profile")
def update_industry_profile_endpoint(profile: str):
    """Update industry profile."""
    from .config_manager import update_industry_profile

    try:
        update_industry_profile(profile)
        return JSONResponse({"message": f"Industry profile updated to {profile}"})
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


# ============ Job Tracking Endpoints ============


class JobStatusUpdateRequest(BaseModel):
    status: str
    notes: str | None = None


class CompanyLinkRequest(BaseModel):
    company_link: str


class CoverLetterRequest(BaseModel):
    """Request model for cover letter generation."""
    job_description: str
    resume_text: str | None = None


@app.get("/api/jobs/tracked")
def get_tracked_jobs(status: str | None = None, include_hidden: bool = False):
    """Get all tracked jobs, optionally filtered by status."""
    tracker = get_tracker()

    # Parse status filter (comma-separated)
    status_filter = None
    if status:
        status_filter = [s.strip() for s in status.split(",") if s.strip()]

    jobs = tracker.get_all_jobs(status_filter=status_filter, include_hidden=include_hidden)
    return JSONResponse({"jobs": jobs, "count": len(jobs)})


@app.get("/api/jobs/{job_id}")
def get_job_details(job_id: str):
    """Get details for a specific job."""
    tracker = get_tracker()
    job = tracker.get_job(job_id)

    if not job:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")

    return JSONResponse({"job": job})


@app.post("/api/jobs/{job_id}/status")
def update_job_status(job_id: str, req: JobStatusUpdateRequest):
    """Update job status and optionally add notes."""
    tracker = get_tracker()

    if req.status not in VALID_STATUSES:
        raise HTTPException(
            status_code=400, detail=f"Invalid status '{req.status}'. Valid statuses: {', '.join(VALID_STATUSES)}"
        )

    success = tracker.update_status(job_id, req.status, req.notes)

    if not success:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")

    job = tracker.get_job(job_id)
    return JSONResponse({"message": "Status updated", "job": job})


@app.post("/api/jobs/{job_id}/hide")
def hide_job(job_id: str):
    """Hide a job (won't appear in search results)."""
    tracker = get_tracker()
    success = tracker.hide_job(job_id)

    if not success:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")

    return JSONResponse({"message": "Job hidden"})


@app.post("/api/jobs/{job_id}/company-link")
def set_company_link(job_id: str, req: CompanyLinkRequest):
    """Set the direct company career page link for a job."""
    tracker = get_tracker()

    # Validate URL format
    if not req.company_link.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="Invalid URL format")

    success = tracker.set_company_link(job_id, req.company_link)

    if not success:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")

    job = tracker.get_job(job_id)
    return JSONResponse({"message": "Company link updated", "job": job})


@app.post("/api/jobs/{job_id}/cover-letter")
def generate_cover_letter(job_id: str, req: CoverLetterRequest):
    """Generate a customized cover letter for a job using Gemini."""
    tracker = get_tracker()
    job = tracker.get_job(job_id)

    if not job:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")

    # Get resume text
    resume_text = req.resume_text
    if not resume_text and RESUME_FILE.exists():
        resume_text = RESUME_FILE.read_text(encoding="utf-8")

    if not resume_text:
        raise HTTPException(status_code=400, detail="No resume text provided")

    # Use provided job description or job summary
    job_description = req.job_description or job.get("summary", "")

    if not job_description:
        raise HTTPException(status_code=400, detail="No job description available")

    # Generate cover letter using Gemini
    try:
        from .gemini_provider import simple_gemini_query

        prompt = f"""Generate a professional cover letter for this job application.

Job Title: {job.get('title', 'N/A')}
Company: {job.get('company', 'N/A')}
Location: {job.get('location', 'N/A')}

Job Description:
{job_description}

Candidate Resume:
{resume_text}

Write a concise, professional cover letter (3-4 paragraphs) that:
1. Expresses interest in the specific role
2. Highlights relevant experience from the resume
3. Explains why the candidate is a good fit
4. Ends with a call to action

Return ONLY the cover letter text, no additional commentary."""

        cover_letter = simple_gemini_query(prompt)

        if not cover_letter:
            raise HTTPException(status_code=500, detail="Failed to generate cover letter")

        return JSONResponse(
            {"cover_letter": cover_letter, "job_title": job.get("title"), "company": job.get("company")}
        )

    except Exception as e:
        print(f"Cover letter generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate cover letter: {str(e)}") from e


@app.post("/api/jobs/find-company-link/{job_id}")
async def find_company_link(job_id: str):
    """Find direct company career page link for an aggregator job.

    Extracts company name from job and searches for direct company careers page.
    """
    from .mcp_providers import generate_job_leads_via_mcp

    tracker = get_tracker()
    job = tracker.get_job(job_id)

    if not job:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")

    company = job.get("company", "")
    title = job.get("title", "")

    if not company:
        raise HTTPException(status_code=400, detail="Job has no company name")

    # Search for this specific company's jobs using CompanyJobs
    try:
        search_query = f"{title} at {company}"
        print(f"Finding company link: searching for '{search_query}'")

        # Use CompanyJobs provider to find company career page
        leads = generate_job_leads_via_mcp(
            query=search_query, count=5, count_per_provider=5, location=job.get("location", "Remote")
        )

        # Filter results to only this company's jobs (case-insensitive)
        company_lower = company.lower()
        company_jobs = [
            lead
            for lead in leads
            if lead.get("company", "").lower() == company_lower
            and lead.get("source") == "CompanyJobs"  # Only direct company links
        ]

        if not company_jobs:
            return JSONResponse(
                {
                    "found": False,
                    "message": f"No direct company links found for {company}. Try searching their website directly.",
                }
            )

        # Return the first match (most relevant)
        direct_link = company_jobs[0].get("link", "")

        # Automatically save it to the job
        if direct_link:
            tracker.set_company_link(job_id, direct_link)

        return JSONResponse(
            {
                "found": True,
                "company_link": direct_link,
                "job": company_jobs[0],
                "message": f"Found direct link at {company}",
            }
        )

    except Exception as e:
        print(f"Error finding company link: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to find company link: {str(e)}") from e
